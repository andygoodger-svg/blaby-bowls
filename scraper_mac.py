#!/usr/bin/env python3
"""
Blaby Bowls Scraper v3
- Hinckley: per-team fixture URLs (no duplicates)
- Leicester: Division 1 ONLY (docx download)
- South Leics: Google Sheets CSV export (no JS needed)
- Telegram notifications on completion
"""
import requests
from bs4 import BeautifulSoup
import os, subprocess, sys, re, csv, io
from datetime import datetime

OUTPUT_DIR = "/Users/andrewgoodger/blaby-bowls"
HEADERS = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) BlabyScraper/3.0"}


def _load_dotenv(path):
    """Tiny .env loader: KEY=VALUE per line, # comments, optional quotes.
    Only sets vars that aren't already in the environment."""
    try:
        with open(path) as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, _, val = line.partition("=")
                key = key.strip()
                val = val.strip().strip('"').strip("'")
                os.environ.setdefault(key, val)
    except FileNotFoundError:
        pass


_load_dotenv(os.path.join(OUTPUT_DIR, ".env"))

# Telegram notification — configure via environment variables or .env file.
# See .env.example for the expected keys.
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID", "")


def send_telegram(message):
    """Send a notification to Telegram."""
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        print("  [INFO] Telegram not configured (set TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID in .env) - skipping.")
        return
    try:
        url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
        r = requests.post(url, json={"chat_id": TELEGRAM_CHAT_ID, "text": message, "parse_mode": "HTML"}, timeout=10)
        if r.status_code == 200:
            print("  Telegram notification sent.")
        else:
            print(f"  [WARN] Telegram send failed: {r.status_code}")
    except Exception as e:
        print(f"  [WARN] Telegram error: {e}")


# =========================================================================
# HINCKLEY - Per-team fixture pages (clean, no duplicates)
# =========================================================================
HINCKLEY_BASE = "https://www.bowlsresultstwo.co.uk/results24"
HINCKLEY_TEAMS = [
    {"name": "Blaby A", "div": 1, "div_name": "Division 1", "team_id": 2},
    {"name": "Blaby B", "div": 1, "div_name": "Division 1", "team_id": 6},
    {"name": "Blaby C", "div": 4, "div_name": "Division 4", "team_id": 33},
]

# =========================================================================
# LEICESTER - Division 1 ONLY
# =========================================================================
LEICESTER_DOCX = {
    "div1": {
        "name": "Division 1",
        "url": "https://www.leicesterbowlsleague.co.uk/shared/attachments.asp?f=0eec92c2%2D60ad%2D405e%2Db88d%2D59e311d126fd%2Edocx&o=Division%2D1%2DResults%2DFixtures%2Dand%2Dresults%2D2026%2Edocx"
    },
    "tables": {
        "name": "League Tables 2026",
        "url": "https://www.leicesterbowlsleague.co.uk/shared/attachments.asp?f=da5c79df%2D2126%2D472c%2D928d%2Ddd8f4e9999d3%2Edocx&o=League%2DTables%2D2026%2Edocx"
    }
}

# =========================================================================
# SOUTH LEICS - Google Sheets IDs (extracted from embedded iframes)
# =========================================================================
# League/Fixtures page has 6 sheets (one per division)
SOUTH_LEICS_FIXTURES_SHEETS = [
    {"id": "1NIqYReT3YvUXoX8hNnqKnpkaEYNfha8ifOdHWyM2R_w", "name": "Div 1 Fixtures"},
    {"id": "1XaP3Uv2l4pR71KwdXMNmdPknsvCX3wX0IZHDN5IaXfM", "name": "Div 2 Fixtures"},
    {"id": "10J734xth5HL_Y_LTqkXCc6d1yJC1R1nhcKWK863_tIk", "name": "Div 3 Fixtures"},
    {"id": "1-lOR1raGH6mcoZ0XQVKtuYGSXrwB68uMkH7xpG_7jf4", "name": "Div 4 Fixtures"},
    {"id": "1eJS6sDVrdRL2ZKdrnCnzUCWdzbcyCl-NHTyt7-5oI78", "name": "Div 5 Fixtures"},
    {"id": "1BYZA0ktPqoJCxf5P76m_H7mIYP9RDr_PAEEvcOxe8Qo", "name": "Div 6 Fixtures"},
]
SOUTH_LEICS_RESULTS_SHEET = "1-FoyAaFLqhC7POnhBP-KY7agQ5WXxEXe3WGB5rwpqXw"
SOUTH_LEICS_TABLES_SHEET = "1b8vRpHme_v5oBJqJ9R3eHP5jQOoeLSbZObRmME7OYtI"


CSS = """<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:transparent;color:#333;padding:10px}
h2{color:#1a5c2e;margin:15px 0 8px;font-size:1.2em;border-bottom:2px solid #1a5c2e;padding-bottom:4px}
h3{color:#2d7a45;margin:10px 0 6px;font-size:1em}
table{border-collapse:collapse;width:100%;margin-bottom:15px;font-size:.85em}
th{background:#1a5c2e;color:#fff;padding:6px 8px;text-align:left;font-weight:600}
td{padding:5px 8px;border-bottom:1px solid #e0e0e0}
tr:nth-child(even){background:#f5f9f6}
tr.blaby-row{background:#d4edda!important;font-weight:600}
tr.blaby-match{background:#e8f5e9}
.updated{font-size:.75em;color:#888;margin-top:10px;text-align:right}
.no-data{color:#888;font-style:italic;padding:10px}
.fixture-date{background:#1a5c2e;color:#fff;padding:4px 8px;font-weight:600;font-size:.9em}
.score{font-weight:700;color:#1a5c2e}
.league-header{background:#2d7a45;color:#fff;padding:8px 12px;font-size:1.05em;margin-top:20px;border-radius:4px}
.team-header{background:#1a5c2e;color:#fff;padding:6px 10px;font-weight:700;font-size:.9em;border-radius:3px 3px 0 0;margin-top:12px}
</style>"""

def html_wrap(title, body):
    now = datetime.now().strftime("%d %b %Y %H:%M")
    return f'<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0"><title>{title}</title>{CSS}</head><body>{body}<p class="updated">Last updated: {now}</p></body></html>'


def fetch(url):
    try:
        r = requests.get(url, headers=HEADERS, timeout=30)
        r.raise_for_status()
        return BeautifulSoup(r.text, "html.parser")
    except Exception as e:
        print(f"  [ERROR] {url}: {e}")
        return None


def fetch_google_sheet_csv(sheet_id, label="sheet"):
    """Fetch a Google Sheet as CSV. Returns list of rows (each row is a list of strings)."""
    # Try CSV export first, then fall back to htmlembed + parsing
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid=0"
    print(f"  Fetching Google Sheet CSV: {label}")
    try:
        r = requests.get(url, timeout=30)
        if r.status_code == 200:
            reader = csv.reader(io.StringIO(r.text))
            rows = [row for row in reader]
            print(f"    Got {len(rows)} rows")
            return rows
        else:
            print(f"    [WARN] CSV export returned {r.status_code}, trying htmlembed...")
            return fetch_google_sheet_html(sheet_id, label)
    except Exception as e:
        print(f"    [ERROR] {e}")
        return []


def fetch_google_sheet_html(sheet_id, label="sheet"):
    """Fallback: try gviz/tq endpoint, then htmlembed with table parsing."""
    # Try gviz/tq endpoint (sometimes works when export doesn't)
    gviz_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv"
    try:
        r = requests.get(gviz_url, timeout=30)
        if r.status_code == 200 and len(r.text) > 50:
            reader = csv.reader(io.StringIO(r.text))
            rows = [row for row in reader]
            if rows:
                print(f"    Got {len(rows)} rows via gviz/tq")
                return rows
    except Exception as e:
        print(f"    [WARN] gviz/tq failed: {e}")

    # Try htmlembed
    url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/htmlembed"
    try:
        r = requests.get(url, timeout=30)
        if r.status_code == 200:
            soup = BeautifulSoup(r.text, "html.parser")
            rows = []
            for table in soup.find_all("table"):
                for tr in table.find_all("tr"):
                    cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                    if any(c for c in cells):
                        rows.append(cells)
            print(f"    Got {len(rows)} rows via htmlembed")
            return rows
        else:
            print(f"    [WARN] htmlembed also failed: {r.status_code}")
            return []
    except Exception as e:
        print(f"    [ERROR] htmlembed: {e}")
        return []


# =========================================================================
# HINCKLEY - Per-team fixture pages (clean, no duplicates)
# =========================================================================
def scrape_hinckley_team_fixtures(team):
    """Scrape a single team's fixture page. Returns list of dicts."""
    url = f"{HINCKLEY_BASE}/teamfixtures.php?ly=0&f=0&yearid=2026&d={team['div']}&t={team['team_id']}&web=hinckley&m=0&leagueid=1"
    print(f"  {team['name']} fixtures: {url}")
    soup = fetch(url)
    if not soup:
        return []

    fixtures = []
    tables = soup.find_all("table")
    for table in tables:
        rows = table.find_all("tr")
        for row in rows:
            cells = [c.get_text(strip=True) for c in row.find_all("td")]
            if len(cells) >= 3:
                date = cells[0]
                opponent = cells[1]
                home_away = cells[2] if len(cells) > 2 else ""

                if date and opponent and any(m in date for m in ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]):
                    is_home = "Home" in home_away or "home" in home_away
                    fixtures.append({
                        "date": date,
                        "opponent": opponent,
                        "home_away": "Home" if is_home else "Away",
                        "home": team["name"] if is_home else opponent,
                        "away": opponent if is_home else team["name"],
                        "score": None,
                        "team": team["name"]
                    })

                # Check for second fixture in same row (pages show 2 columns)
                if len(cells) >= 6:
                    date2 = cells[3]
                    opponent2 = cells[4]
                    home_away2 = cells[5] if len(cells) > 5 else ""
                    if date2 and opponent2 and any(m in date2 for m in ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]):
                        is_home2 = "Home" in home_away2 or "home" in home_away2
                        fixtures.append({
                            "date": date2,
                            "opponent": opponent2,
                            "home_away": "Home" if is_home2 else "Away",
                            "home": team["name"] if is_home2 else opponent2,
                            "away": opponent2 if is_home2 else team["name"],
                            "score": None,
                            "team": team["name"]
                        })

    return fixtures


def scrape_hinckley_table(div_id):
    """Scrape league table for a division."""
    url = f"{HINCKLEY_BASE}/tables.php?f=0&d={div_id}&yearid=2026&web=hinckley&m=0&leagueid=1"
    print(f"  Hinckley Div {div_id} table: {url}")
    soup = fetch(url)
    if not soup:
        return '<p class="no-data">No league table data available yet.</p>'

    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if len(rows) > 2:
            html = str(table)
            ns = BeautifulSoup(html, "html.parser")
            for row in ns.find_all("tr"):
                if "blaby" in row.get_text().lower():
                    row["class"] = row.get("class", []) + ["blaby-row"]
            return str(ns)

    return '<p class="no-data">No league table data available yet.</p>'


# =========================================================================
# LEICESTER - Division 1 only
# =========================================================================
def parse_leicester_docx(url, label):
    """Download a .docx and extract data. Returns structured fixture/result data."""
    try:
        from docx import Document
        import tempfile
        print(f"  Downloading: {label}")
        r = requests.get(url, headers=HEADERS, timeout=30)
        r.raise_for_status()
        tmp = os.path.join(tempfile.gettempdir(), "leic_bowls.docx")
        with open(tmp, "wb") as f:
            f.write(r.content)
        doc = Document(tmp)

        # Extract ALL table rows (for full table display) and Blaby rows
        blaby_rows = []
        all_rows = []
        tables_raw = []  # Keep raw table structures for better parsing

        for t_idx, t in enumerate(doc.tables):
            table_rows = []
            header = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                row_data = {"cells": cells, "header": header, "table_idx": t_idx}
                all_rows.append(row_data)
                table_rows.append(cells)
                if "blaby" in " ".join(cells).lower():
                    blaby_rows.append(row_data)
            tables_raw.append({"header": header, "rows": table_rows})

        # Extract all paragraphs (for fixture dates context)
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        blaby_paras = [p for p in all_paras if "blaby" in p.lower()]

        # DEBUG: Print raw table structure so we can understand the format
        print(f"    DEBUG: {len(doc.tables)} table(s), {len(all_rows)} total rows, {len(blaby_rows)} Blaby rows")
        print(f"    DEBUG: {len(all_paras)} paragraphs")
        for t_idx, t_data in enumerate(tables_raw):
            ncols = len(t_data['header']) if t_data['header'] else (len(t_data['rows'][0]) if t_data['rows'] else 0)
            print(f"    DEBUG Table {t_idx}: {ncols} cols, header={t_data['header']}, {len(t_data['rows'])} rows")
            for r_idx, row in enumerate(t_data['rows'][:16]):  # Show first 16 rows, ALL columns
                print(f"      Row {r_idx}: {row}")
            if len(t_data['rows']) > 16:
                print(f"      ... ({len(t_data['rows'])-16} more rows)")
        # Show first few paragraphs for context
        for p in all_paras[:5]:
            print(f"    DEBUG Para: '{p[:80]}'")

        os.remove(tmp)
        return {
            "rows": blaby_rows,
            "all_rows": all_rows,
            "paragraphs": blaby_paras,
            "all_paras": all_paras,
            "tables_raw": tables_raw
        }
    except ImportError:
        print("  [WARN] python-docx not installed")
        return None
    except Exception as e:
        print(f"  [ERROR] docx parse: {e}")
        return None


def parse_leicester_fixtures_structured(div1_data):
    """Parse Leicester Div 1 docx into structured fixture list with dates.

    The docx has a WIDE table with TWO side-by-side columns of fixture boxes:
      - Left column group:  date in col 0, teams in col 2, shots in col 3
      - Right column group: date in col N, teams in col N+2, shots in col N+3
    where N is the first column of the right-side boxes.

    We auto-detect column groups by scanning every column in the first few rows
    for date-day patterns (e.g. "29th", "6th"), then process each group
    independently using the same logic.
    """
    if not div1_data:
        return []

    tables_raw = div1_data.get("tables_raw", [])
    fixtures = []

    MONTHS = ["january", "february", "march", "april", "may", "june",
              "july", "august", "september", "october", "november", "december"]
    date_day_pattern = re.compile(r'^(\d{1,2})(st|nd|rd|th)$', re.IGNORECASE)
    SKIP_WORDS = {"shots", "w", "l", "d", "for", "agst", "home", "away", "pts",
                  "points", "rinks", "total"}

    for t_data in tables_raw:
        rows = t_data["rows"]
        if not rows:
            continue

        ncols = max(len(r) for r in rows)
        print(f"    Leicester table: {len(rows)} rows x {ncols} cols")

        # --- Auto-detect column groups ---
        # Scan all rows for columns that contain date-day patterns (e.g. "29th")
        date_cols = set()
        for row in rows:
            for ci, cell in enumerate(row):
                val = cell.strip()
                if date_day_pattern.match(val):
                    date_cols.add(ci)

        if not date_cols:
            print(f"    [WARN] No date columns found in Leicester table")
            continue

        # Group nearby date columns (they should be the start of each fixture block)
        # Sort and cluster — columns within 3 of each other are the same group
        sorted_cols = sorted(date_cols)
        groups = []
        for col in sorted_cols:
            if not groups or col - groups[-1][-1] > 5:
                groups.append([col])
            else:
                groups[-1].append(col)

        # For each group, use the minimum column index as the base
        group_bases = [min(g) for g in groups]
        print(f"    Detected {len(group_bases)} column group(s) starting at columns: {group_bases}")

        # --- Process each column group ---
        for base_col in group_bases:
            date_col = base_col       # Date day/month
            team_col = base_col + 2   # Team names
            shots_col = base_col + 3  # Shots/scores

            current_date_day = ""
            current_date_full = ""
            team_buffer = []

            for row in rows:
                col_date = row[date_col].strip() if len(row) > date_col else ""
                col_team = row[team_col].strip() if len(row) > team_col else ""
                col_shots = row[shots_col].strip() if len(row) > shots_col else ""

                # Check for date day (e.g. "29th")
                if date_day_pattern.match(col_date):
                    if team_buffer:
                        _flush_team_pairs(team_buffer, current_date_full, fixtures, col3_list=[])
                        team_buffer = []
                    current_date_day = col_date
                    current_date_full = col_date
                    continue

                # Check for month name
                if col_date.lower() in MONTHS:
                    current_date_full = f"{current_date_day} {col_date}" if current_date_day else col_date
                    continue

                # Skip header/label rows
                if col_date.lower() in SKIP_WORDS or col_team.lower() in SKIP_WORDS:
                    continue

                # Team name
                if col_team and col_team.lower() not in SKIP_WORDS:
                    score_info = col_shots if col_shots and any(c.isdigit() for c in col_shots) else ""
                    team_buffer.append({"name": col_team, "score": score_info, "date": current_date_full})

            # Flush remaining pairs for this group
            if team_buffer:
                _flush_team_pairs(team_buffer, current_date_full, fixtures, col3_list=[])

    # Sort fixtures by date for consistent display
    MONTH_ORDER = {m: i for i, m in enumerate(MONTHS)}
    def date_sort_key(f):
        parts = f.get("date", "").split()
        if len(parts) >= 2:
            day_match = date_day_pattern.match(parts[0])
            day = int(day_match.group(1)) if day_match else 99
            month = MONTH_ORDER.get(parts[1].lower(), 99)
            return (month, day)
        return (99, 99)

    fixtures.sort(key=date_sort_key)

    print(f"    Leicester parsed: {len(fixtures)} Blaby fixtures found")
    for f in fixtures:
        status = f"({f['score']})" if f.get('score') else "(upcoming)"
        print(f"      {f['date']}: {f['home']} v {f['away']} {status}")

    return fixtures


def _flush_team_pairs(team_buffer, default_date, fixtures, col3_list):
    """Pair up consecutive teams from the buffer as Home/Away fixtures."""
    for j in range(0, len(team_buffer) - 1, 2):
        home_entry = team_buffer[j]
        away_entry = team_buffer[j + 1]
        home = home_entry["name"]
        away = away_entry["name"]
        date = home_entry.get("date", default_date)

        # Determine score (combine home and away shots if both present)
        score = None
        h_score = home_entry.get("score", "")
        a_score = away_entry.get("score", "")
        if h_score and a_score:
            score = f"{h_score} - {a_score}"
        elif h_score or a_score:
            score = h_score or a_score

        if "blaby" in home.lower() or "blaby" in away.lower():
            fixtures.append({
                "date": date,
                "home": home,
                "away": away,
                "score": score
            })


def parse_leicester_table_div1(tables_data):
    """Extract Division 1 section from the league tables docx."""
    if not tables_data or not tables_data["all_rows"]:
        return '<p class="no-data">No Leicester league table data yet.</p>'

    # The tables docx has all 3 divisions in sequence.
    # We need to find the Division 1 section and stop before Division 2.
    all_rows = tables_data["all_rows"]
    div1_rows = []
    in_div1 = False

    for row_data in all_rows:
        cells = row_data["cells"]
        text = " ".join(cells).strip().lower()

        # Detect division headers
        if "division 1" in text and "division 2" not in text:
            in_div1 = True
            continue
        elif "division 2" in text:
            in_div1 = False
            continue

        if in_div1:
            # Skip empty rows
            if not any(c.strip() for c in cells):
                continue
            div1_rows.append(cells)

    if not div1_rows:
        # Fallback: if no division headers found, just show all rows
        div1_rows = [r["cells"] for r in all_rows if any(c.strip() for c in r["cells"])]

    if not div1_rows:
        return '<p class="no-data">No Division 1 league table data yet.</p>'

    # Build HTML table
    html = '<table>\n'
    for i, cells in enumerate(div1_rows):
        is_header = i == 0 and any(h in " ".join(cells).lower() for h in ["pld", "played", "won", "pts", "team"])
        is_blaby = "blaby" in " ".join(cells).lower()

        if is_header:
            html += '<tr>' + ''.join(f'<th>{c}</th>' for c in cells[:8] if c.strip()) + '</tr>\n'
        else:
            cls = ' class="blaby-row"' if is_blaby else ""
            html += f'<tr{cls}>' + ''.join(f'<td>{c}</td>' for c in cells[:8] if c.strip() or i > 0) + '</tr>\n'
    html += '</table>\n'
    return html


# =========================================================================
# SOUTH LEICESTERSHIRE TRIPLES - Google Sheets CSV
# =========================================================================
def scrape_south_leics():
    """Scrape South Leics via Google Sheets CSV export."""
    results = {"fixtures": [], "results": [], "tables": "", "fixture_divs": {}}

    # --- FIXTURES: Check all 6 division sheets for Blaby teams ---
    print("  Fetching fixture sheets...")
    for sheet in SOUTH_LEICS_FIXTURES_SHEETS:
        rows = fetch_google_sheet_csv(sheet["id"], sheet["name"])
        if not rows:
            continue

        # Find Blaby fixtures in this division
        blaby_fixtures = []
        current_date = ""
        div_name = sheet["name"].replace(" Fixtures", "")

        for row in rows:
            if not any(cell.strip() for cell in row):
                continue

            # Check if this row is a date header
            first_cell = row[0].strip() if row else ""
            if re.match(r'\d{1,2}(st|nd|rd|th)\s+(January|February|March|April|May|June|July|August|September|October|November|December)', first_cell, re.IGNORECASE):
                current_date = first_cell
                continue

            # Check if row mentions Blaby
            row_text = " ".join(row).lower()
            if "blaby" in row_text:
                # Fixture format: Home team | [score] | Away team | [score]
                # or just: Home team | Away team
                clean = [c.strip() for c in row if c.strip()]
                if len(clean) >= 2:
                    blaby_fixtures.append({
                        "date": current_date,
                        "cells": clean,
                        "div": div_name
                    })

        if blaby_fixtures:
            results["fixture_divs"][div_name] = blaby_fixtures
            results["fixtures"].extend(blaby_fixtures)
            print(f"    {div_name}: {len(blaby_fixtures)} Blaby fixtures found")
        else:
            print(f"    {div_name}: no Blaby fixtures")

    # --- RESULTS ---
    print("  Fetching results sheet...")
    rows = fetch_google_sheet_csv(SOUTH_LEICS_RESULTS_SHEET, "Match Results")
    if rows:
        current_date = ""
        current_section = ""
        for row in rows:
            if not any(cell.strip() for cell in row):
                continue
            first_cell = row[0].strip() if row else ""

            # Check for date headers
            if re.match(r'\d{1,2}(st|nd|rd|th)\s+\w+\s+\d{4}', first_cell):
                current_date = first_cell
                continue

            # Check for section headers (e.g. "SEMI FINALS", "QUARTER FINALS")
            if first_cell and first_cell.isupper() and len(first_cell) > 3:
                current_section = first_cell
                continue

            row_text = " ".join(row).lower()
            if "blaby" in row_text:
                clean = [c.strip() for c in row if c.strip()]
                if clean:
                    results["results"].append({
                        "date": current_date,
                        "section": current_section,
                        "cells": clean
                    })

    # --- TABLES ---
    # Skip for now - current sheet is 2025 Partridge Cup data
    # Will re-enable once 2026 league tables are published
    print("  Tables: skipping (2026 tables not yet published, using placeholders)")
    rows = None  # Was: fetch_google_sheet_csv(SOUTH_LEICS_TABLES_SHEET, "League Tables")
    if rows:
        # Parse the table data - find groups/divisions containing Blaby
        current_group = ""
        table_html = ""
        current_table_rows = []
        current_headers = []

        for row in rows:
            if not any(cell.strip() for cell in row):
                # Empty row - might be end of a group
                if current_table_rows:
                    # Check if this group has Blaby
                    group_text = " ".join(str(r) for r in current_table_rows).lower()
                    if "blaby" in group_text:
                        table_html += f'<h3>{current_group}</h3>\n<table>\n'
                        if current_headers:
                            table_html += '<tr>' + ''.join(f'<th>{h}</th>' for h in current_headers) + '</tr>\n'
                        for tr in current_table_rows:
                            is_blaby = "blaby" in " ".join(tr).lower()
                            cls = ' class="blaby-row"' if is_blaby else ""
                            table_html += f'<tr{cls}>' + ''.join(f'<td>{c}</td>' for c in tr) + '</tr>\n'
                        table_html += '</table>\n'
                    current_table_rows = []
                    current_headers = []
                continue

            clean = [c.strip() for c in row]

            # Check for group/division headers
            first = clean[0] if clean else ""
            if re.match(r'(Group|Division)\s+', first, re.IGNORECASE) and len([c for c in clean if c]) <= 2:
                if current_table_rows:
                    group_text = " ".join(str(r) for r in current_table_rows).lower()
                    if "blaby" in group_text:
                        table_html += f'<h3>{current_group}</h3>\n<table>\n'
                        if current_headers:
                            table_html += '<tr>' + ''.join(f'<th>{h}</th>' for h in current_headers) + '</tr>\n'
                        for tr in current_table_rows:
                            is_blaby = "blaby" in " ".join(tr).lower()
                            cls = ' class="blaby-row"' if is_blaby else ""
                            table_html += f'<tr{cls}>' + ''.join(f'<td>{c}</td>' for c in tr) + '</tr>\n'
                        table_html += '</table>\n'
                current_table_rows = []
                current_headers = []
                current_group = first
                continue

            # Check for header rows (Pl., W, D, L, etc.)
            if any(h in " ".join(clean).upper() for h in ["PL.", "PL", "PLAYED", "PTS"]):
                current_headers = [c for c in clean if c]
                continue

            # Regular data row
            if any(c for c in clean if c):
                current_table_rows.append([c for c in clean if c])

        # Flush last group
        if current_table_rows:
            group_text = " ".join(str(r) for r in current_table_rows).lower()
            if "blaby" in group_text:
                table_html += f'<h3>{current_group}</h3>\n<table>\n'
                if current_headers:
                    table_html += '<tr>' + ''.join(f'<th>{h}</th>' for h in current_headers) + '</tr>\n'
                for tr in current_table_rows:
                    is_blaby = "blaby" in " ".join(tr).lower()
                    cls = ' class="blaby-row"' if is_blaby else ""
                    table_html += f'<tr{cls}>' + ''.join(f'<td>{c}</td>' for c in tr) + '</tr>\n'
                table_html += '</table>\n'

        results["tables"] = table_html

    return results


# =========================================================================
# HTML GENERATION
# =========================================================================
def gen_fixture_table(team_label, fixtures, show_date_rows=False):
    """Generate a consistent fixture table for any league.
    fixtures: list of dicts with keys: date, home, away (and optionally score)
    """
    upcoming = [f for f in fixtures if not f.get("score")]
    if not upcoming:
        return f'<p class="no-data">No upcoming {team_label} fixtures found.</p>\n'

    html = f'<div class="team-header">{team_label}</div>\n'
    html += '<table>\n'
    html += '<tr><th>Date</th><th>Home</th><th></th><th>Away</th></tr>\n'

    current_date = ""
    for f in upcoming:
        date_str = f.get("date", "")
        if show_date_rows and date_str and date_str != current_date:
            current_date = date_str
            html += f'<tr><td colspan="4" class="fixture-date">{current_date}</td></tr>\n'
            date_str = ""  # Don't repeat date in the row

        html += f'<tr class="blaby-match"><td>{date_str if not show_date_rows else ""}</td><td>{f["home"]}</td><td>V</td><td>{f["away"]}</td></tr>\n'
    html += '</table>\n'
    return html


def gen_fixtures(hinckley_data, south_leics, leicester_data):
    b = "<h2>Blaby Bowls - Upcoming Fixtures 2026</h2>\n"

    # --- Hinckley ---
    b += '<div class="league-header">Hinckley &amp; District Triples League</div>\n'
    current_div = None
    for team in HINCKLEY_TEAMS:
        if team["div_name"] != current_div:
            current_div = team["div_name"]
            b += f'<h3>{current_div}</h3>\n'

        raw_fixtures = hinckley_data.get(team["name"], [])
        fixtures = [{"date": f["date"], "home": f["home"], "away": f["away"], "score": f["score"]} for f in raw_fixtures]
        b += gen_fixture_table(f'{team["name"]} Fixtures', fixtures)

    # --- South Leics ---
    b += '<div class="league-header">South Leicestershire Triples League</div>\n'
    if south_leics["fixture_divs"]:
        for div_name, raw_fixtures in south_leics["fixture_divs"].items():
            # Determine Blaby team name from div
            team_map = {"Div 1": "Blaby A", "Div 2": "Blaby B", "Div 3": "Blaby C"}
            team_name = team_map.get(div_name, "Blaby")

            fixtures = []
            for f in raw_fixtures:
                cells = f["cells"]
                if len(cells) >= 2:
                    home = cells[0]
                    away = cells[-1] if len(cells) == 2 else cells[1]
                    if len(cells) == 4:
                        home = cells[0]
                        away = cells[2]
                    fixtures.append({"date": f["date"], "home": home, "away": away, "score": None})

            b += f'<h3>{div_name}</h3>\n'
            b += gen_fixture_table(f'{team_name} Fixtures — {div_name}', fixtures)
    else:
        b += '<p class="no-data">No Blaby fixtures found yet. Season starts 28th April 2026.</p>\n'

    # --- Leicester Division 1 ---
    b += '<div class="league-header">Leicester Bowls League — Division 1</div>\n'
    leic_div1 = leicester_data.get("div1")
    leic_fixtures = parse_leicester_fixtures_structured(leic_div1)
    if leic_fixtures:
        upcoming = [f for f in leic_fixtures if not f.get("score")]
        if upcoming:
            b += gen_fixture_table('Blaby Fixtures — Division 1', upcoming)
        else:
            b += '<p class="no-data">No upcoming Blaby Division 1 fixtures.</p>\n'
    elif leic_div1 and leic_div1["rows"]:
        # Fallback: show raw Blaby rows in best format we can
        b += '<div class="team-header">Blaby Fixtures — Division 1</div>\n'
        b += '<table>\n<tr><th>Date</th><th>Home</th><th></th><th>Away</th></tr>\n'

        # Try to pair up raw rows as home/away
        raw_cells = []
        date_str = ""
        for r in leic_div1["rows"]:
            if r.get("header"):
                date_str = " ".join(h for h in r["header"][:2] if h)
            for c in r["cells"]:
                val = c.strip()
                if val:
                    # Split on newlines (cells might contain "Team1\nTeam2")
                    if '\n' in val:
                        raw_cells.extend([p.strip() for p in val.split('\n') if p.strip()])
                    else:
                        raw_cells.append(val)

        # Pair consecutive team names
        for j in range(0, len(raw_cells) - 1, 2):
            home = raw_cells[j]
            away = raw_cells[j + 1]
            b += f'<tr class="blaby-match"><td>{date_str}</td><td>{home}</td><td>V</td><td>{away}</td></tr>\n'
            date_str = ""  # Only show date on first row

        b += '</table>\n'
    else:
        b += '<p class="no-data">No Blaby Division 1 fixtures found yet.</p>\n'

    return html_wrap("Blaby Bowls - Fixtures 2026", b)


def gen_result_table(team_label, results_list):
    """Generate a consistent result table for any league.
    results_list: list of dicts with keys: date, home, away, score (or home_score/away_score)
    """
    if not results_list:
        return f'<p class="no-data">No {team_label} results yet.</p>\n'

    html = f'<div class="team-header">{team_label}</div>\n'
    html += '<table>\n'
    html += '<tr><th>Date</th><th>Home</th><th>Score</th><th>Away</th></tr>\n'
    for r in results_list:
        score = r.get("score", "")
        html += f'<tr class="blaby-match"><td>{r["date"]}</td><td>{r["home"]}</td><td class="score">{score}</td><td>{r["away"]}</td></tr>\n'
    html += '</table>\n'
    return html


def gen_results(hinckley_data, south_leics, leicester_data):
    b = "<h2>Blaby Bowls - Results &amp; Scores 2026</h2>\n"

    # --- Hinckley ---
    b += '<div class="league-header">Hinckley &amp; District Triples League</div>\n'
    current_div = None
    for team in HINCKLEY_TEAMS:
        if team["div_name"] != current_div:
            current_div = team["div_name"]
            b += f'<h3>{current_div}</h3>\n'

        raw_results = [f for f in hinckley_data.get(team["name"], []) if f["score"] is not None]
        results_list = [{"date": f["date"], "home": f["home"], "away": f["away"], "score": f["score"]} for f in raw_results]
        b += gen_result_table(f'{team["name"]} Results', results_list)

    # --- South Leics (2026 only) ---
    b += '<div class="league-header">South Leicestershire Triples League</div>\n'
    results_2026 = [r for r in south_leics["results"] if "2026" in r.get("date", "")]
    if results_2026:
        formatted = []
        for r in results_2026:
            cells = r["cells"]
            date = r["date"]
            if len(cells) >= 4:
                formatted.append({"date": date, "home": cells[0], "away": cells[2], "score": f'{cells[1]} - {cells[3]}'})
            elif len(cells) >= 2:
                formatted.append({"date": date, "home": cells[0], "away": cells[-1], "score": ""})
        b += gen_result_table('South Leics Results', formatted)
    else:
        b += '<p class="no-data">No South Leics results yet. Season starts 28th April 2026.</p>\n'

    # --- Leicester Division 1 ---
    b += '<div class="league-header">Leicester Bowls League — Division 1</div>\n'
    leic_div1 = leicester_data.get("div1")
    leic_fixtures = parse_leicester_fixtures_structured(leic_div1)
    if leic_fixtures:
        played = [f for f in leic_fixtures if f.get("score")]
        b += gen_result_table('Blaby Results — Division 1', played)
    else:
        b += '<p class="no-data">No Division 1 results available yet.</p>\n'

    return html_wrap("Blaby Bowls - Results 2026", b)


def gen_table_page(league, div, html):
    b = f"<h2>{league}</h2>\n<h3>{div}</h3>\n{html}"
    return html_wrap(f"Blaby - {league} {div}", b)


def gen_index():
    b = """<h2>Blaby Bowls Club - 2026 Season</h2>
<p>Auto-updated fixtures, results, and league tables for all Blaby teams.</p>
<h3>Fixtures &amp; Results</h3>
<table><tr><th>Page</th><th>Description</th></tr>
<tr><td><a href="fixtures.html">All Fixtures</a></td><td>Upcoming matches for all Blaby teams</td></tr>
<tr><td><a href="results.html">All Results</a></td><td>Match results and scores as they come in</td></tr></table>
<h3>League Tables</h3>
<table><tr><th>League</th><th>Division</th></tr>
<tr><td><a href="table-hinckley-div1.html">Hinckley &amp; District Triples</a></td><td>Division 1 (Blaby A, Blaby B)</td></tr>
<tr><td><a href="table-hinckley-div4.html">Hinckley &amp; District Triples</a></td><td>Division 4 (Blaby C)</td></tr>
<tr><td><a href="table-south-leics.html">South Leicestershire Triples</a></td><td>League Table</td></tr>
<tr><td><a href="table-leicester.html">Leicester Bowls League</a></td><td>Division 1</td></tr></table>
<h3>Source Websites</h3>
<table><tr><th>League</th><th>Link</th></tr>
<tr><td>Hinckley &amp; District Triples</td><td><a href="https://www.bowlsresultstwo.co.uk/hinckley/" target="_blank">bowlsresultstwo.co.uk</a></td></tr>
<tr><td>South Leicestershire Triples</td><td><a href="https://sites.google.com/view/southleicestershiretriples/" target="_blank">Google Sites</a></td></tr>
<tr><td>Leicester Bowls League</td><td><a href="https://www.leicesterbowlsleague.co.uk/community/leicester-bowls-league-15027/" target="_blank">leicesterbowlsleague.co.uk</a></td></tr></table>"""
    return html_wrap("Blaby Bowls Club - 2026 Season", b)


def git_push():
    try:
        os.chdir(OUTPUT_DIR)
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        subprocess.run(["git", "add", "-A"], check=True, capture_output=True)
        r = subprocess.run(["git", "status", "--porcelain"], capture_output=True, text=True)
        if not r.stdout.strip():
            print("  No changes to push.")
            return True
        subprocess.run(["git", "commit", "-m", f"Auto-update: {now}"], check=True, capture_output=True)
        subprocess.run(["git", "push"], check=True, capture_output=True)
        print(f"  Pushed to GitHub at {now}")
        return True
    except subprocess.CalledProcessError as e:
        print(f"  [ERROR] Git push failed: {e}")
        return False


def main():
    print("=" * 60)
    print(f"Blaby Bowls Scraper v3 - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("=" * 60)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # --- HINCKLEY ---
    print("\n[1/3] Hinckley & District Triples League...")
    hinckley_data = {}
    for team in HINCKLEY_TEAMS:
        hinckley_data[team["name"]] = scrape_hinckley_team_fixtures(team)
        print(f"    {team['name']}: {len(hinckley_data[team['name']])} fixtures")

    hinckley_tables = {}
    for div_id in set(t["div"] for t in HINCKLEY_TEAMS):
        hinckley_tables[div_id] = scrape_hinckley_table(div_id)

    # --- SOUTH LEICS ---
    print("\n[2/3] South Leicestershire Triples League...")
    south_leics = scrape_south_leics()
    sl_fix = sum(len(v) for v in south_leics["fixture_divs"].values())
    print(f"    Total: {sl_fix} fixtures, {len(south_leics['results'])} results")

    # --- LEICESTER (Division 1 only) ---
    print("\n[3/3] Leicester Bowls League (Division 1 only)...")
    leicester_data = {}
    for key, info in LEICESTER_DOCX.items():
        leicester_data[key] = parse_leicester_docx(info["url"], info["name"])
        if leicester_data[key]:
            print(f"    {info['name']}: {len(leicester_data[key]['rows'])} Blaby rows, {len(leicester_data[key]['all_rows'])} total rows")
        else:
            print(f"    {info['name']}: failed or no data")

    # --- GENERATE HTML ---
    print("\nGenerating HTML files...")

    # Leicester table - Division 1 only
    leic_table_html = parse_leicester_table_div1(leicester_data.get("tables"))

    # South Leics table — placeholders for 2026 (tables not yet published)
    south_leics_table = '<h3>League Tables 2026</h3>\n'
    south_leics_table += '<p class="no-data">League tables will appear here once the 2026 season is underway (first matches 28th April). '
    south_leics_table += 'Check back after the first round of games.</p>\n'
    south_leics_table += '<h3>Partridge Cup 2026</h3>\n'
    south_leics_table += '<p class="no-data">Partridge Cup tables will appear here once the 2026 competition starts.</p>\n'

    files = {
        "index.html": gen_index(),
        "fixtures.html": gen_fixtures(hinckley_data, south_leics, leicester_data),
        "results.html": gen_results(hinckley_data, south_leics, leicester_data),
        "table-hinckley-div1.html": gen_table_page("Hinckley & District Triples", "Division 1", hinckley_tables.get(1, '<p class="no-data">No data yet.</p>')),
        "table-hinckley-div4.html": gen_table_page("Hinckley & District Triples", "Division 4", hinckley_tables.get(4, '<p class="no-data">No data yet.</p>')),
        "table-south-leics.html": gen_table_page("South Leicestershire Triples", "League Table", south_leics_table),
        "table-leicester.html": gen_table_page("Leicester Bowls League", "Division 1", leic_table_html),
    }

    for fn, content in files.items():
        path = os.path.join(OUTPUT_DIR, fn)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"  Written: {fn}")

    # --- GIT PUSH ---
    print("\nPushing to GitHub Pages...")
    pushed = git_push()

    # --- TELEGRAM NOTIFICATION ---
    now = datetime.now().strftime("%d %b %Y %H:%M")
    h_total = sum(len(hinckley_data.get(t["name"], [])) for t in HINCKLEY_TEAMS)
    l_total = len(leicester_data.get("div1", {}).get("rows", []) if leicester_data.get("div1") else [])

    if pushed:
        msg = (
            f"<b>Blaby Bowls Updated</b> - {now}\n\n"
            f"Hinckley: {h_total} fixtures\n"
            f"Leicester Div 1: {l_total} Blaby rows\n"
            f"South Leics: {sl_fix} fixtures, {len(south_leics['results'])} results\n\n"
            f"<a href='https://andygoodger-svg.github.io/blaby-bowls/'>View site</a>"
        )
        send_telegram(msg)
        print("\nDone! Live at https://andygoodger-svg.github.io/blaby-bowls/")
    else:
        send_telegram(f"Blaby Bowls scraper ran at {now} but push failed. Check logs.")
        print("\n[WARN] Git push failed - files saved locally only.")

    print("=" * 60)
    return True


if __name__ == "__main__":
    main()
