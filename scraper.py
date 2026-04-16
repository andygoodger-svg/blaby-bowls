#!/usr/bin/env python3
"""
Blaby Bowls Scraper v2
Uses per-team fixture URLs for Hinckley (no duplicates),
proper docx download URLs for Leicester,
and notes South Leics JS limitation.
"""
import requests
from bs4 import BeautifulSoup
import os, subprocess, sys, re
from datetime import datetime

OUTPUT_DIR = "/opt/blaby-bowls"
HEADERS = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) BlabyScraper/2.0"}

# Hinckley team-specific URLs (no more division mixing)
HINCKLEY_BASE = "https://www.bowlsresultstwo.co.uk/results24"
HINCKLEY_TEAMS = [
    {"name": "Blaby A", "div": 1, "div_name": "Division 1", "team_id": 2},
    {"name": "Blaby B", "div": 1, "div_name": "Division 1", "team_id": 6},
    {"name": "Blaby C", "div": 4, "div_name": "Division 4", "team_id": 33},
]

# Leicester docx URLs (use attachments.asp pattern)
LEICESTER_DOCX = {
    "div1": {
        "name": "Division 1",
        "url": "https://www.leicesterbowlsleague.co.uk/shared/attachments.asp?f=0eec92c2%2D60ad%2D405e%2Db88d%2D59e311d126fd%2Edocx&o=Division%2D1%2DResults%2DFixtures%2Dand%2Dresults%2D2026%2Edocx"
    },
    "div2n": {
        "name": "Division 2 North",
        "url": "https://www.leicesterbowlsleague.co.uk/shared/attachments.asp?f=f916e73d%2D27eb%2D456e%2Da68e%2D7a4dcbf4292d%2Edocx&o=Division%2D2%2DNorth%2DFixtures%2Dand%2DResults%2D2026%2Edocx"
    },
    "div2s": {
        "name": "Division 2 South",
        "url": "https://www.leicesterbowlsleague.co.uk/shared/attachments.asp?f=8870c9b7%2D6e1a%2D4eca%2Db528%2Dc9d5832f007a%2Edocx&o=Division%2D2%2DSouth%2DFixtures%2Dand%2DResults%2D2026%2Edocx"
    },
    "tables": {
        "name": "League Tables 2026",
        "url": "https://www.leicesterbowlsleague.co.uk/shared/attachments.asp?f=da5c79df%2D2126%2D472c%2D928d%2Ddd8f4e9999d3%2Edocx&o=League%2DTables%2D2026%2Edocx"
    }
}

CSS = """<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:transparent;color:#333;padding:10px}
h2{color:#1a5c2e;margin:15px 0 8px;font-size:1.2em;border-bottom:2px solid #1a5c2e;padding-bottom:4px}
h3{color:#2d7a45;margin:10px 0 6px;font-size:1em}
table{border-collapse:collapse;width:100%;margin-bottom:15px;font-size:.85em}
th{background:#1a5c2e;color:#fff;padding:6px 8px;text-align:left;font-weight:600}
td{padding:5px 8px;border-bottom:1px solid #e0e0e0}
tr:nth-child(even){background:#f5f9f6}
tr.blaby-row{background:#d4edda!important;font-weight:600}
tr.blaby-match{background:#e8f5e9}
.updated{font-size:.75em;color:#888;margin-top:10px;text-align:right}
.no-data{color:#888;font-style:italic;padding:10px}
.fixture-date{background:#1a5c2e;color:#fff;padding:4px 8px;font-weight:600;font-size:.9em}
.score{font-weight:700;color:#1a5c2e}
.league-header{background:#2d7a45;color:#fff;padding:8px;font-size:1em;margin-top:15px}
</style>"""


def html_wrap(title, body):
    now = datetime.now().strftime("%d %b %Y %H:%M")
    return f'<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0"><title>{title}</title>{CSS}</head><body>{body}<p class="updated">Last updated: {now}</p></body></html>'


def fetch(url):
    try:
        r = requests.get(url, headers=HEADERS, timeout=30)
        r.raise_for_status()
        return BeautifulSoup(r.text, "html.parser")
    except Exception as e:
        print(f"  [ERROR] {url}: {e}")
        return None


# =========================================================================
# HINCKLEY - Per-team fixture pages (clean, no duplicates)
# =========================================================================
def scrape_hinckley_team_fixtures(team):
    """Scrape a single team's fixture page. Returns list of dicts."""
    url = f"{HINCKLEY_BASE}/teamfixtures.php?ly=0&f=0&yearid=2026&d={team['div']}&t={team['team_id']}&web=hinckley&m=0&leagueid=1"
    print(f"  {team['name']} fixtures: {url}")
    soup = fetch(url)
    if not soup:
        return []

    fixtures = []
    tables = soup.find_all("table")
    for table in tables:
        rows = table.find_all("tr")
        for row in rows:
            cells = [c.get_text(strip=True) for c in row.find_all("td")]
            if len(cells) >= 3:
                # Format: Date | Opponent | Home/Away | [Score fields if played]
                # or: Date | Opponent | Home/Away | Date | Opponent | Home/Away
                # Team fixture pages show: date, opponent, H/A
                date = cells[0]
                opponent = cells[1]
                home_away = cells[2] if len(cells) > 2 else ""

                if date and opponent and any(m in date for m in ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]):
                    is_home = "Home" in home_away or "home" in home_away
                    fixtures.append({
                        "date": date,
                        "opponent": opponent,
                        "home_away": "Home" if is_home else "Away",
                        "home": team["name"] if is_home else opponent,
                        "away": opponent if is_home else team["name"],
                        "score": None,
                        "team": team["name"]
                    })

                # Check for second fixture in same row (pages show 2 columns)
                if len(cells) >= 6:
                    date2 = cells[3]
                    opponent2 = cells[4]
                    home_away2 = cells[5] if len(cells) > 5 else ""
                    if date2 and opponent2 and any(m in date2 for m in ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]):
                        is_home2 = "Home" in home_away2 or "home" in home_away2
                        fixtures.append({
                            "date": date2,
                            "opponent": opponent2,
                            "home_away": "Home" if is_home2 else "Away",
                            "home": team["name"] if is_home2 else opponent2,
                            "away": opponent2 if is_home2 else team["name"],
                            "score": None,
                            "team": team["name"]
                        })

    return fixtures


def scrape_hinckley_table(div_id):
    """Scrape league table for a division."""
    url = f"{HINCKLEY_BASE}/tables.php?res=1&d=0&yearid=2026&web=hinckley&leagueid={div_id}"
    print(f"  Hinckley Div {div_id} table: {url}")
    soup = fetch(url)
    if not soup:
        return '<p class="no-data">No league table data available yet.</p>'

    # Find tables with data
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if len(rows) > 2:
            html = str(table)
            # Highlight Blaby rows
            ns = BeautifulSoup(html, "html.parser")
            for row in ns.find_all("tr"):
                if "blaby" in row.get_text().lower():
                    row["class"] = row.get("class", []) + ["blaby-row"]
            return str(ns)

    return '<p class="no-data">No league table data available yet.</p>'


# =========================================================================
# LEICESTER - Download and parse .docx files
# =========================================================================
def parse_leicester_docx(url, label):
    """Download a .docx and extract Blaby-related data."""
    try:
        from docx import Document
        import tempfile
        print(f"  Downloading: {label}")
        r = requests.get(url, headers=HEADERS, timeout=30)
        r.raise_for_status()
        tmp = os.path.join(tempfile.gettempdir(), "leic_bowls.docx")
        with open(tmp, "wb") as f:
            f.write(r.content)
        doc = Document(tmp)

        # Extract table rows containing "blaby"
        blaby_rows = []
        all_rows = []
        for t in doc.tables:
            header = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                all_rows.append(cells)
                if "blaby" in " ".join(cells).lower():
                    blaby_rows.append({"cells": cells, "header": header})

        # Also get paragraphs
        blaby_paras = [p.text.strip() for p in doc.paragraphs if "blaby" in p.text.lower()]

        os.remove(tmp)
        return {"rows": blaby_rows, "all_rows": all_rows, "paragraphs": blaby_paras}
    except ImportError:
        print("  [WARN] python-docx not installed")
        return None
    except Exception as e:
        print(f"  [ERROR] docx parse: {e}")
        return None


# =========================================================================
# SOUTH LEICESTERSHIRE TRIPLES
# =========================================================================
def scrape_south_leics():
    """
    South Leics uses Google Sites which renders via JavaScript.
    Plain requests can't see the embedded Google Sheets content.
    We try anyway in case it works, but this likely returns empty.
    """
    results = {"fixtures": [], "results": [], "tables": ""}

    for page in ["league", "results", "tables"]:
        url = f"https://sites.google.com/view/southleicestershiretriples/{page}"
        print(f"  South Leics {page}: {url}")
        soup = fetch(url)
        if not soup:
            continue

        # Try to find embedded Google Sheets iframes
        iframes = soup.find_all("iframe")
        for iframe in iframes:
            src = iframe.get("src", "")
            if "docs.google.com" in src or "sheet" in src.lower():
                print(f"    Found sheet iframe: {src[:80]}...")
                sheet_soup = fetch(src)
                if sheet_soup:
                    for table in sheet_soup.find_all("table"):
                        if "blaby" in table.get_text().lower():
                            for row in table.find_all("tr"):
                                if "blaby" in row.get_text().lower():
                                    cells = [c.get_text(strip=True) for c in row.find_all(["td", "th"])]
                                    if page == "league":
                                        results["fixtures"].append(cells)
                                    elif page == "results":
                                        results["results"].append(cells)

                            if page == "tables":
                                ns = BeautifulSoup(str(table), "html.parser")
                                for row in ns.find_all("tr"):
                                    if "blaby" in row.get_text().lower():
                                        row["class"] = row.get("class", []) + ["blaby-row"]
                                results["tables"] += str(ns)

        # Also check direct table content
        for table in soup.find_all("table"):
            if "blaby" in table.get_text().lower():
                for row in table.find_all("tr"):
                    if "blaby" in row.get_text().lower():
                        cells = [c.get_text(strip=True) for c in row.find_all(["td", "th"])]
                        if page in ["league"]:
                            results["fixtures"].append(cells)
                        elif page == "results":
                            results["results"].append(cells)

    return results


# =========================================================================
# HTML GENERATION
# =========================================================================
def gen_fixtures(hinckley_data, south_leics, leicester_data):
    b = "<h2>Blaby Bowls - Upcoming Fixtures 2026</h2>\n"

    # Hinckley
    b += '<div class="league-header">Hinckley &amp; District Triples League</div>\n'
    current_div = None
    for team in HINCKLEY_TEAMS:
        if team["div_name"] != current_div:
            current_div = team["div_name"]
            b += f'<h3>{current_div}</h3>\n'

        fixtures = [f for f in hinckley_data.get(team["name"], []) if f["score"] is None]
        if fixtures:
            b += f'<table><tr><th colspan="4">{team["name"]} Fixtures</th></tr>\n'
            b += '<tr><th>Date</th><th>Home</th><th></th><th>Away</th></tr>\n'
            for f in fixtures:
                b += f'<tr class="blaby-match"><td>{f["date"]}</td><td>{f["home"]}</td><td>V</td><td>{f["away"]}</td></tr>\n'
            b += '</table>\n'
        else:
            b += f'<p class="no-data">No upcoming {team["name"]} fixtures found.</p>\n'

    # South Leics
    b += '<div class="league-header">South Leicestershire Triples League</div>\n'
    if south_leics["fixtures"]:
        b += '<table><tr><th>Fixture Details</th></tr>\n'
        for f in south_leics["fixtures"]:
            b += f'<tr class="blaby-match"><td>{" | ".join(f)}</td></tr>\n'
        b += '</table>\n'
    else:
        b += '<p class="no-data">Season starts 28th April 2026. Fixtures will appear once available.</p>\n'

    # Leicester
    b += '<div class="league-header">Leicester Bowls League</div>\n'
    has_leic = False
    for key in ["div1", "div2n", "div2s"]:
        data = leicester_data.get(key)
        if data and data["rows"]:
            # Filter for fixture-like rows (those without scores/digits)
            fixture_rows = [r for r in data["rows"] if not any(any(c.isdigit() for c in cell) for cell in r["cells"][1:])]
            if fixture_rows:
                b += f'<h3>{LEICESTER_DOCX[key]["name"]}</h3>\n'
                b += '<table><tr><th>' + '</th><th>'.join(fixture_rows[0].get("header", ["Details"])[:6]) + '</th></tr>\n'
                for r in fixture_rows:
                    b += f'<tr class="blaby-match"><td>' + '</td><td>'.join(r["cells"][:6]) + '</td></tr>\n'
                b += '</table>\n'
                has_leic = True
    if not has_leic:
        leic_found = any(leicester_data.get(k, {}).get("rows") for k in ["div1", "div2n", "div2s"])
        if leic_found:
            # Show whatever we found
            for key in ["div1", "div2n", "div2s"]:
                data = leicester_data.get(key)
                if data and data["rows"]:
                    b += f'<h3>{LEICESTER_DOCX[key]["name"]}</h3>\n'
                    b += '<table>\n'
                    for r in data["rows"]:
                        b += f'<tr class="blaby-match"><td>' + '</td><td>'.join(r["cells"][:6]) + '</td></tr>\n'
                    b += '</table>\n'
                    has_leic = True
    if not has_leic:
        b += '<p class="no-data">No Blaby fixtures found in Leicester Bowls League documents yet.</p>\n'

    return html_wrap("Blaby Bowls - Fixtures 2026", b)


def gen_results(hinckley_data, south_leics, leicester_data):
    b = "<h2>Blaby Bowls - Results &amp; Scores 2026</h2>\n"

    # Hinckley
    b += '<div class="league-header">Hinckley &amp; District Triples League</div>\n'
    current_div = None
    for team in HINCKLEY_TEAMS:
        if team["div_name"] != current_div:
            current_div = team["div_name"]
            b += f'<h3>{current_div}</h3>\n'

        results = [f for f in hinckley_data.get(team["name"], []) if f["score"] is not None]
        if results:
            b += f'<table><tr><th colspan="4">{team["name"]} Results</th></tr>\n'
            b += '<tr><th>Date</th><th>Home</th><th>Score</th><th>Away</th></tr>\n'
            for f in results:
                b += f'<tr class="blaby-match"><td>{f["date"]}</td><td>{f["home"]}</td><td class="score">{f["score"]}</td><td>{f["away"]}</td></tr>\n'
            b += '</table>\n'
        else:
            b += f'<p class="no-data">No {team["name"]} results yet - season starts May 2026.</p>\n'

    # South Leics
    b += '<div class="league-header">South Leicestershire Triples League</div>\n'
    if south_leics["results"]:
        b += '<table><tr><th>Result Details</th></tr>\n'
        for r in south_leics["results"]:
            b += f'<tr class="blaby-match"><td>{" | ".join(r)}</td></tr>\n'
        b += '</table>\n'
    else:
        b += '<p class="no-data">Season starts 28th April 2026. Results will appear once matches are played.</p>\n'

    # Leicester
    b += '<div class="league-header">Leicester Bowls League</div>\n'
    has_leic = False
    for key in ["div1", "div2n", "div2s"]:
        data = leicester_data.get(key)
        if data and data["rows"]:
            result_rows = [r for r in data["rows"] if any(any(c.isdigit() for c in cell) for cell in r["cells"][1:])]
            if result_rows:
                b += f'<h3>{LEICESTER_DOCX[key]["name"]}</h3>\n'
                b += '<table>\n'
                if result_rows[0].get("header"):
                    b += '<tr><th>' + '</th><th>'.join(result_rows[0]["header"][:6]) + '</th></tr>\n'
                for r in result_rows:
                    b += f'<tr class="blaby-match"><td>' + '</td><td>'.join(r["cells"][:6]) + '</td></tr>\n'
                b += '</table>\n'
                has_leic = True
    if not has_leic:
        b += '<p class="no-data">No results available yet for Leicester Bowls League.</p>\n'

    return html_wrap("Blaby Bowls - Results 2026", b)


def gen_table_page(league, div, html):
    b = f"<h2>{league}</h2>\n<h3>{div}</h3>\n{html}"
    return html_wrap(f"Blaby - {league} {div}", b)


def gen_index():
    b = """<h2>Blaby Bowls Club - 2026 Season</h2>
<p>Auto-updated fixtures, results, and league tables for all Blaby teams.</p>
<h3>Fixtures &amp; Results</h3>
<table><tr><th>Page</th><th>Description</th></tr>
<tr><td><a href="fixtures.html">All Fixtures</a></td><td>Upcoming matches for all Blaby teams</td></tr>
<tr><td><a href="results.html">All Results</a></td><td>Match results and scores as they come in</td></tr></table>
<h3>League Tables</h3>
<table><tr><th>League</th><th>Division</th></tr>
<tr><td><a href="table-hinckley-div1.html">Hinckley &amp; District Triples</a></td><td>Division 1 (Blaby A, Blaby B)</td></tr>
<tr><td><a href="table-hinckley-div4.html">Hinckley &amp; District Triples</a></td><td>Division 4 (Blaby C)</td></tr>
<tr><td><a href="table-south-leics.html">South Leicestershire Triples</a></td><td>League Table</td></tr>
<tr><td><a href="table-leicester.html">Leicester Bowls League</a></td><td>League Table</td></tr></table>
<h3>Source Websites</h3>
<table><tr><th>League</th><th>Link</th></tr>
<tr><td>Hinckley &amp; District Triples</td><td><a href="https://www.bowlsresultstwo.co.uk/hinckley/" target="_blank">bowlsresultstwo.co.uk</a></td></tr>
<tr><td>South Leicestershire Triples</td><td><a href="https://sites.google.com/view/southleicestershiretriples/" target="_blank">Google Sites</a></td></tr>
<tr><td>Leicester Bowls League</td><td><a href="https://www.leicesterbowlsleague.co.uk/community/leicester-bowls-league-15027/" target="_blank">leicesterbowlsleague.co.uk</a></td></tr></table>"""
    return html_wrap("Blaby Bowls Club - 2026 Season", b)


def git_push():
    try:
        os.chdir(OUTPUT_DIR)
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        subprocess.run(["git", "add", "-A"], check=True, capture_output=True)
        r = subprocess.run(["git", "status", "--porcelain"], capture_output=True, text=True)
        if not r.stdout.strip():
            print("  No changes to push.")
            return True
        subprocess.run(["git", "commit", "-m", f"Auto-update: {now}"], check=True, capture_output=True)
        subprocess.run(["git", "push"], check=True, capture_output=True)
        print(f"  Pushed to GitHub at {now}")
        return True
    except subprocess.CalledProcessError as e:
        print(f"  [ERROR] Git push failed: {e}")
        return False


def main():
    print("=" * 60)
    print(f"Blaby Bowls Scraper v2 - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("=" * 60)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # --- HINCKLEY ---
    print("\n[1/3] Hinckley & District Triples League...")
    hinckley_data = {}
    for team in HINCKLEY_TEAMS:
        hinckley_data[team["name"]] = scrape_hinckley_team_fixtures(team)
        print(f"    {team['name']}: {len(hinckley_data[team['name']])} fixtures")

    hinckley_tables = {}
    for div_id in set(t["div"] for t in HINCKLEY_TEAMS):
        hinckley_tables[div_id] = scrape_hinckley_table(div_id)

    # --- SOUTH LEICS ---
    print("\n[2/3] South Leicestershire Triples League...")
    south_leics = scrape_south_leics()
    print(f"    {len(south_leics['fixtures'])} fixtures, {len(south_leics['results'])} results")

    # --- LEICESTER ---
    print("\n[3/3] Leicester Bowls League...")
    leicester_data = {}
    for key, info in LEICESTER_DOCX.items():
        leicester_data[key] = parse_leicester_docx(info["url"], info["name"])
        if leicester_data[key]:
            print(f"    {info['name']}: {len(leicester_data[key]['rows'])} Blaby rows")
        else:
            print(f"    {info['name']}: failed or no data")

    # --- GENERATE HTML ---
    print("\nGenerating HTML files...")

    # Leicester table content
    leic_table_html = '<p class="no-data">No Leicester league table data yet.</p>'
    leic_tables = leicester_data.get("tables")
    if leic_tables and leic_tables["rows"]:
        leic_table_html = '<table>\n'
        if leic_tables["rows"][0].get("header"):
            leic_table_html += '<tr><th>' + '</th><th>'.join(leic_tables["rows"][0]["header"][:8]) + '</th></tr>\n'
        for r in leic_tables["rows"]:
            leic_table_html += '<tr class="blaby-row"><td>' + '</td><td>'.join(r["cells"][:8]) + '</td></tr>\n'
        leic_table_html += '</table>\n'
        # Also show full table context if available
        if leic_tables["all_rows"]:
            leic_table_html = '<table>\n'
            for r in leic_tables["all_rows"]:
                cls = ' class="blaby-row"' if "blaby" in " ".join(r).lower() else ""
                leic_table_html += f'<tr{cls}><td>' + '</td><td>'.join(r[:8]) + '</td></tr>\n'
            leic_table_html += '</table>\n'

    south_leics_table = south_leics["tables"] if south_leics["tables"] else '<p class="no-data">South Leics table data is embedded in Google Sheets and not available via automated scraping. Check the website directly.</p>'

    files = {
        "index.html": gen_index(),
        "fixtures.html": gen_fixtures(hinckley_data, south_leics, leicester_data),
        "results.html": gen_results(hinckley_data, south_leics, leicester_data),
        "table-hinckley-div1.html": gen_table_page("Hinckley & District Triples", "Division 1", hinckley_tables.get(1, '<p class="no-data">No data yet.</p>')),
        "table-hinckley-div4.html": gen_table_page("Hinckley & District Triples", "Division 4", hinckley_tables.get(4, '<p class="no-data">No data yet.</p>')),
        "table-south-leics.html": gen_table_page("South Leicestershire Triples", "League Table", south_leics_table),
        "table-leicester.html": gen_table_page("Leicester Bowls League", "League Table", leic_table_html),
    }

    for fn, content in files.items():
        path = os.path.join(OUTPUT_DIR, fn)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"  Written: {fn}")

    # --- GIT PUSH ---
    print("\nPushing to GitHub Pages...")
    if git_push():
        print("\nDone! Live at https://andygoodger-svg.github.io/blaby-bowls/")
    else:
        print("\n[WARN] Git push failed - files saved locally only.")

    print("=" * 60)
    return True


if __name__ == "__main__":
    main()
